# Text extraction (PDF, TXT, etc.)
"""
Document processing module with multi-format support.
Handles PDF, TXT, DOCX, and CSV files with LangChain integration.
"""
import logging
from pathlib import Path
from typing import List
from docx import Document
import pandas as pd
from langchain_community.document_loaders import (
    UnstructuredPDFLoader,  # Robust PDF loader
    UnstructuredFileLoader   # Universal text loader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

# Initialize logger for error tracking
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Core processor for document loading and text splitting"""
    
    def __init__(self, chunk_size: int = 7500, chunk_overlap: int = 100):
        """
        Initialize text splitter with configuration
        
        Args:
            chunk_size: Character length for text chunks
            chunk_overlap: Overlap between chunks for context preservation
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    def load_document(self, file_path: Path) -> List:
        """
        Main document loading router
        
        Args:
            file_path: Path to target file
            
        Returns:
            List of LangChain Document objects
            
        Raises:
            ValueError: For unsupported file types
            Exception: Detailed error logging
        """
        try:
            logger.info(f"Processing file: {file_path}")
            
            # Route to appropriate loader based on extension
            if file_path.suffix.lower() == ".pdf":
                return self._load_pdf(file_path)
            elif file_path.suffix.lower() == ".txt":
                return self._load_txt(file_path)
            elif file_path.suffix.lower() == ".docx":
                return self._load_docx(file_path)
            elif file_path.suffix.lower() == ".csv":
                return self._load_csv(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Document loading failed: {e}")
            raise
    
    # --------------------------------------------------
    # EXPLICIT LOADER IMPLEMENTATIONS
    # --------------------------------------------------
    
    def _load_pdf(self, file_path: Path) -> List:
        """
        PDF loader using UnstructuredPDFLoader
        
        Advantages:
        - Preserves complex layouts
        - Handles embedded tables/images
        - Better than PyPDF2 for structured PDFs
        """
        loader = UnstructuredPDFLoader(str(file_path))
        return loader.load()
    
    def _load_txt(self, file_path: Path) -> List:
        """
        Plain text loader using UnstructuredFileLoader
        
        Features:
        - Automatic encoding detection
        - Preserves original line breaks
        - Minimal preprocessing
        """
        loader = UnstructuredFileLoader(str(file_path))
        return loader.load()
    
    def _load_docx(self, file_path: Path) -> List:
        """
        DOCX loader using python-docx.
        
        Note:
        - Requires python-docx package
        - Only extracts plain text (no styling)
        """
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [LangchainDocument(page_content=text)]
    
    def _load_csv(self, file_path: Path) -> List:
        """
        CSV loader using pandas.
        
        Options:
        - Use CSVLoader for simple cases
        - pandas provides better table handling
        """
        df = pd.read_csv(file_path)
        text = df.to_string(index=False)
        return [LangchainDocument(page_content=text)]
    
    def split_documents(self, documents: List) -> List:
        """
        Split documents using recursive text splitting.
        
        Process:
        1. Splits on paragraphs
        2. Then sentences
        3. Finally by chunk size
        
        Returns:
            List of processed document chunks
        """
        try:
            logger.info(f"Splitting {len(documents)} documents")
            return self.splitter.split_documents(documents)
        except Exception as e:
            logger.error(f"Text splitting failed: {e}")
            raise
        