import streamlit as st
import base64
import os
import pandas as pd
from docx import Document
from core.auth import database, models

def display_file_viewer(temp_dir):
    """Display a preview for files PDF, DOCX, CSV, TXT"""
    if not st.session_state.get("uploaded_files"):
        return
    
    try:
        first_file = st.session_state.uploaded_files[0]
        file_path = os.path.join(temp_dir, first_file.name)
        file_extension = os.path.splitext(file_path)[1].lower()

        # PDF
        if file_extension == ".pdf":
            with open(file_path, "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode("utf-8")

            st.markdown(
                f"""
                <iframe src="data:application/pdf;base64,{base64_pdf}"
                        width="100%" height="600" type="application/pdf"
                        style="border-radius: 10px;">
                </iframe>
                """,
                unsafe_allow_html=True,
            )

        # DOCX
        elif file_extension == ".docx":
            doc = Document(file_path)
            content = [p.text for p in doc.paragraphs if p.text.strip()]

            st.markdown(
                f"""
                <div style="border:1px solid #ddd; border-radius:10px; padding:10px;
                            max-height:600px; overflow-y:auto; background-color:#f9f9f9;">
                    {'<br>'.join(content)}
                </div>
                """,
                unsafe_allow_html=True,
            )

        # CSV
        elif file_extension == ".csv":
            df = pd.read_csv(file_path)
            st.dataframe(df, height=400, use_container_width=True)

        # TXT
        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            st.markdown(
                f"""
                <div style="border:1px solid #ddd; border-radius:10px; padding:10px;
                            max-height:600px; overflow-y:auto; background-color:#f9f9f9;
                            white-space: pre-wrap; font-family: monospace;">
                    {content}
                </div>
                """,
                unsafe_allow_html=True,
            )

        else:
            st.warning("Format non supporté (PDF, DOCX, CSV, TXT uniquement).")

    except Exception as e:
        st.warning(f"Visualisation impossible : {str(e)}")

def display_file_viewer_by_id(doc_id: int):
    """Display a preview of the file stored in the database via its id"""
    db = database.SessionLocal()
    doc = db.get(models.Document, doc_id)
    if not doc:
        st.error("Document introuvable.")
        return

    try:
        file_extension = os.path.splitext(doc.path)[1].lower()

        # PDF
        if file_extension == ".pdf":
            with open(doc.path, "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode("utf-8")

            st.markdown(
                f"""
                <iframe src="data:application/pdf;base64,{base64_pdf}"
                        width="100%" height="600" type="application/pdf"
                        style="border-radius: 10px;">
                </iframe>
                """,
                unsafe_allow_html=True,
            )

        # DOCX
        elif file_extension == ".docx":
            document = Document(doc.path)
            content = [p.text for p in document.paragraphs if p.text.strip()]

            st.markdown(
                f"""
                <div style="border:1px solid #ddd; border-radius:10px; padding:10px;
                            max-height:600px; overflow-y:auto; background-color:#f9f9f9;">
                    {'<br>'.join(content)}
                </div>
                """,
                unsafe_allow_html=True,
            )

        # CSV
        elif file_extension == ".csv":
            df = pd.read_csv(doc.path)
            st.dataframe(df, height=400, use_container_width=True)

        # TXT
        elif file_extension == ".txt":
            with open(doc.path, "r", encoding="utf-8") as f:
                content = f.read()

            st.markdown(
                f"""
                <div style="border:1px solid #ddd; border-radius:10px; padding:10px;
                            max-height:600px; overflow-y:auto; background-color:#f9f9f9;
                            white-space: pre-wrap; font-family: monospace;">
                    {content}
                </div>
                """,
                unsafe_allow_html=True,
            )

        else:
            st.warning("Format non supporté (PDF, DOCX, CSV, TXT uniquement).")

    except Exception as exc:
        st.warning(f"Visualisation impossible : {exc}")

# Alias for compatibility
display_pdf_viewer = display_file_viewer
display_pdf_viewer_by_id = display_file_viewer_by_id